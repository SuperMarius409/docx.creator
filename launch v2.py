import re
import docx
import wikipedia
from docx import Document
from urllib.request import urlopen
from bs4 import BeautifulSoup
from requests import get 
import urllib.request
import wikipedia
import requests
import time
import glob
import shutil
import os


#settings

name = input("Enter your name: ")
language = input("The language (en): ")
wikipedia.set_lang(language)
title = input("What do you want your project to be about?\n")
while True:
    try:
        wiki = wikipedia.page(title)
        break
    except:
        print("Project name invalid")
        title = input("Enter another project name: \n")

#downloader

link = (wikipedia.page(title).url)
print(link)
html = urlopen(link)
bs = BeautifulSoup(html, 'html.parser')
images = bs.find_all('img', {'src':re.compile('.*?')})
f= open("cache.txt","w+")
for image in images: 
    url = ('https:' + image['src']+'\n')
    f.write(url)
with open('cache.txt') as f:
   for line in f:
      url = line
      path = 'image'+url.split('/', -1)[-1]
      urllib.request.urlretrieve(url, path.rstrip('\n'))
time.sleep(5)
src_dir = "C:/Users/MMA/OneDrive/Desktop/VSFiles"
dst_dir = "C:/Users/MMA/OneDrive/Desktop/VSFiles/images"
for jpgfile in glob.iglob(os.path.join(src_dir, "*.jpg")):
    shutil.move(jpgfile, dst_dir)

path = os.chdir("C:\\Users\\MMA\\OneDrive\\Desktop\\VSFiles\\images")
i = 0
for file in os.listdir(path):
    new_file_name = "image {}.jpg".format(i)
    os.rename(file, new_file_name)
    i = i + 1

#wikipedia

os.chdir("C:\\Users\\MMA\\OneDrive\\Desktop\\VSFiles")
text = wiki.content
text = re.sub(r'==', '', text)
text = re.sub(r'=', '', text)
text = re.sub(r'\n', '\n    ', text)
split = text.split('Vezi și', 1)
text = split[0]
print('You can close this window')

#docx

document = Document()

paragraph = document.add_heading(title, 0)
paragraph.alignment = 1
paragraph = document.add_paragraph('    ' + text)
paragraph = document.add_paragraph(name)
paragraph.alignment = 2
document.save(title + ".docx")
input()
